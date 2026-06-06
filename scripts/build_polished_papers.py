from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BUNDLED_SOFFICE = Path(
    "/Users/mauriciogarcia/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/soffice"
)
POLISHED_DIR = ROOT / "docs" / "polished"
PUBLIC_POLISHED_DIR = ROOT / "site" / "public" / "paper-assets" / "polished"


@dataclass(frozen=True)
class PolishedPaper:
    slug: str
    title: str
    role: str
    source: Path


PAPERS = [
    PolishedPaper(
        slug="final-report",
        title="Multi-Agent Learning and Imperfect Information Games",
        role="Polished public final report",
        source=POLISHED_DIR / "final-report" / "source.md",
    ),
    PolishedPaper(
        slug="applied-focus",
        title="Applied Focus: Imperfect-Information AI Systems",
        role="Polished public applied-focus paper",
        source=POLISHED_DIR / "applied-focus" / "source.md",
    ),
    PolishedPaper(
        slug="research-notebook",
        title="Research Notebook: Multi-Agent Learning and Liar's Dice",
        role="Polished public research notebook",
        source=POLISHED_DIR / "research-notebook" / "source.md",
    ),
]


def run(command: list[str]) -> None:
    subprocess.run(command, check=True)


def find_soffice() -> Path:
    candidates = [
        Path(value)
        for value in [
            os.environ.get("SOFFICE"),
            str(BUNDLED_SOFFICE) if BUNDLED_SOFFICE.exists() else None,
            shutil.which("soffice"),
            shutil.which("libreoffice"),
        ]
        if value
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("LibreOffice not found. Set SOFFICE=/path/to/soffice and rerun.")


def set_style_font(style, *, size: int, color: str | None = None, bold: bool | None = None) -> None:
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold


def create_reference_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title_style = styles["Title"]
    set_style_font(title_style, size=22, color="191713", bold=True)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(8)

    subtitle_style = styles["Subtitle"]
    set_style_font(subtitle_style, size=12, color="5F594D", bold=False)
    subtitle_style.paragraph_format.space_after = Pt(12)

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        set_style_font(style, size=11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Liar's Dice CFR Lab")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_docx(paper: PolishedPaper, out_dir: Path, reference_docx: Path) -> Path:
    docx_path = out_dir / f"{paper.slug}.docx"
    run(
        [
            "pandoc",
            str(paper.source),
            "--from",
            "markdown",
            "--to",
            "docx",
            "--standalone",
            "--reference-doc",
            str(reference_docx),
            "--metadata",
            f"title={paper.title}",
            "--output",
            str(docx_path),
        ]
    )
    return docx_path


def build_latex(paper: PolishedPaper, out_dir: Path) -> Path:
    tex_path = out_dir / f"{paper.slug}.tex"
    run(
        [
            "pandoc",
            str(paper.source),
            "--from",
            "markdown",
            "--to",
            "latex",
            "--standalone",
            "--metadata",
            f"title={paper.title}",
            "--output",
            str(tex_path),
        ]
    )
    return tex_path


def build_pdf(docx_path: Path, out_dir: Path) -> Path:
    soffice = find_soffice()
    with tempfile.TemporaryDirectory(prefix="liarsdice-polished-soffice-") as profile:
        run(
            [
                str(soffice),
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ]
        )
    return out_dir / f"{docx_path.stem}.pdf"


def copy_public_assets(paper: PolishedPaper, pdf_path: Path, tex_path: Path) -> None:
    PUBLIC_POLISHED_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(pdf_path, PUBLIC_POLISHED_DIR / f"{paper.slug}.pdf")
    public_source_dir = PUBLIC_POLISHED_DIR / paper.slug
    if public_source_dir.exists():
        shutil.rmtree(public_source_dir)
    public_source_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(tex_path, public_source_dir / tex_path.name)


def write_index(papers: list[PolishedPaper]) -> None:
    index = [
        {
            "slug": paper.slug,
            "title": paper.title,
            "role": paper.role,
            "pdf": f"{paper.slug}.pdf",
            "latex": f"{paper.slug}/{paper.slug}.tex",
        }
        for paper in papers
    ]
    (POLISHED_DIR / "index.json").write_text(json.dumps(index, indent=2) + "\n", encoding="utf-8")
    PUBLIC_POLISHED_DIR.mkdir(parents=True, exist_ok=True)
    (PUBLIC_POLISHED_DIR / "index.json").write_text(json.dumps(index, indent=2) + "\n", encoding="utf-8")


def build_polished_papers() -> None:
    if PUBLIC_POLISHED_DIR.exists():
        shutil.rmtree(PUBLIC_POLISHED_DIR)
    with tempfile.TemporaryDirectory(prefix="liarsdice-polished-reference-") as temp_dir:
        reference_docx = Path(temp_dir) / "reference.docx"
        create_reference_docx(reference_docx)
        for paper in PAPERS:
            if not paper.source.exists():
                raise FileNotFoundError(f"Missing polished paper source: {paper.source}")
            out_dir = POLISHED_DIR / paper.slug
            out_dir.mkdir(parents=True, exist_ok=True)
            docx_path = build_docx(paper, out_dir, reference_docx)
            tex_path = build_latex(paper, out_dir)
            pdf_path = build_pdf(docx_path, out_dir)
            copy_public_assets(paper, pdf_path, tex_path)
    write_index(PAPERS)


def main() -> None:
    build_polished_papers()


if __name__ == "__main__":
    main()
