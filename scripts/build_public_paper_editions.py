from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
BUNDLED_SOFFICE = Path(
    "/Users/mauriciogarcia/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/soffice"
)
POLISHED_DIR = ROOT / "docs" / "polished"
PUBLIC_DIR = ROOT / "site" / "public" / "paper-assets" / "polished"
DOWNLOADS = Path.home() / "Downloads"


@dataclass(frozen=True)
class Paper:
    slug: str
    title: str
    role: str
    source: Path
    mode: str


PAPERS = [
    Paper(
        slug="final-report",
        title="Multi-Agent Learning and Imperfect Information Games: Final Report",
        role="Primary final report",
        source=DOWNLOADS / "Final Report.docx",
        mode="report",
    ),
    Paper(
        slug="applied-focus",
        title="Applied Multi-Agent Learning and Imperfect Information Games",
        role="Applied focus deliverable",
        source=DOWNLOADS / "Applied Focus - 1st Deliverable.docx",
        mode="report",
    ),
    Paper(
        slug="research-notebook",
        title="Research Notebook: Multi-Agent Learning and Liar's Dice",
        role="Cleaned research notebook",
        source=DOWNLOADS / "CS370_ Multiagent Learning & Computational Game Solving Notes & Planning.docx",
        mode="notebook",
    ),
]


KNOWN_HEADINGS = {
    "ABSTRACT",
    "INTRODUCTION",
    "HISTORICAL CONTEXT - DEEPER DIVE",
    "KEY FOUNDATIONAL CONCEPTS USED TO BUILD LIBRATUS & PLURIBUS:",
    "VON NEUMANN'S MINIMAX THEOREM:",
    "COUNTERFACTUAL REGRET MINIMIZATION (CFR):",
    "LIBRATUS:",
    "PLURIBUS:",
    "FUTURE WORK DIRECTIONS:",
    "CONCLUSION",
    "REFERENCES:",
    "Background: CFR and Kuhn Poker",
    "Game Design: A Custom 3-Stage Liar's Dice Variant",
    "Game Design: A Custom 3-Stage Liar’s Dice Variant",
    "CFR Training with a Unified Engine",
    "Evaluating Performance and Scalability",
    "Key Lessons Learned",
    "References",
    "Conclusion",
}


NOTEBOOK_REPLACEMENTS = {
    "Due Date: ASK ABOUT THIS, for now, end of the semester": "Due Date: end of semester",
    "To fill. Could use this info to plan ahead.": "Planning note: use this section to organize the final proposal direction.",
    "Texas hold’em rules": "Texas Hold'em Rules Reference",
    "Tab 4": "Mario Kart Research Direction",
    "Tab 5": "Pluribus to Pokemon Battles",
    "Tab 6": "CFR Web Simulation Notes",
    "Tab 7": "",
    "envt": "environment",
}


def run(command: list[str]) -> None:
    subprocess.run(command, check=True)


def find_soffice() -> Path:
    for value in [os.environ.get("SOFFICE"), str(BUNDLED_SOFFICE), shutil.which("soffice"), shutil.which("libreoffice")]:
        if value and Path(value).exists():
            return Path(value)
    raise FileNotFoundError("LibreOffice not found")


def set_style_font(style, name: str, size_pt: float, bold: bool = False, color: str = "191713") -> None:
    style.font.name = name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def apply_style_system(document: Document, mode: str) -> None:
    normal = document.styles["Normal"]
    set_style_font(normal, "Aptos", 10.4 if mode == "report" else 10.2)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color in [
        ("Title", 19, "191713"),
        ("Heading 1", 14, "1F4E79"),
        ("Heading 2", 12, "1F4E79"),
        ("Heading 3", 10.6, "191713"),
    ]:
        if style_name in document.styles:
            set_style_font(document.styles[style_name], "Aptos Display" if style_name == "Title" else "Aptos", size, True, color)
            document.styles[style_name].paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
            document.styles[style_name].paragraph_format.space_after = Pt(5)

    for section in document.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82 if mode == "report" else 0.78)
        section.right_margin = Inches(0.82 if mode == "report" else 0.78)
        section.start_type = WD_SECTION.CONTINUOUS if mode == "notebook" else section.start_type


def classify_paragraph(text: str, index: int, mode: str) -> str | None:
    stripped = text.strip()
    if not stripped:
        return None
    if index < 24 and mode == "report":
        return "title-page"
    if stripped in KNOWN_HEADINGS:
        return "Heading 1"
    if mode == "notebook" and (stripped.startswith("Tab ") or stripped.endswith("Computational Game Solving")):
        return "Heading 1"
    if stripped.endswith(":") and len(stripped) <= 85 and not stripped.startswith("http"):
        return "Heading 2"
    if stripped.isupper() and 3 <= len(stripped) <= 85:
        return "Heading 1"
    return None


def replace_paragraph_text(paragraph, replacements: dict[str, str]) -> None:
    if not replacements:
        return
    text = paragraph.text
    updated = text
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated == text:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)


def remove_empty_notebook_markers(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip() == "":
            continue
        if paragraph.text.strip() in {"Tab 7"}:
            element = paragraph._element
            element.getparent().remove(element)


def polish_paragraphs(document: Document, mode: str) -> None:
    if mode == "notebook":
        remove_empty_notebook_markers(document)
    replacements = NOTEBOOK_REPLACEMENTS if mode == "notebook" else {}
    for index, paragraph in enumerate(document.paragraphs):
        replace_paragraph_text(paragraph, replacements)
        text = paragraph.text.strip()
        classification = classify_paragraph(text, index, mode)

        paragraph.paragraph_format.line_spacing = 1.07
        paragraph.paragraph_format.space_after = Pt(4)

        if classification == "title-page":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if text and index in {9, 10, 15}:
                paragraph.paragraph_format.space_after = Pt(8)
            for run in paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(14 if index in {9, 10, 15} else 10.5)
                run.font.bold = index in {9, 10, 15}
            continue

        if classification in {"Heading 1", "Heading 2"}:
            paragraph.style = document.styles[classification]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.keep_with_next = True
        elif text and mode == "notebook" and paragraph.style.name == "Title":
            paragraph.style = document.styles["Heading 1"]

        for run in paragraph.runs:
            if run.text.strip():
                run.font.name = "Aptos"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
                if run.font.size is None:
                    run.font.size = Pt(10.4 if mode == "report" else 10.2)


def remove_notebook_section_breaks(docx_path: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="paper-ooxml-") as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(docx_path) as zin:
            zin.extractall(tmp_path)

        document_xml = tmp_path / "word" / "document.xml"
        tree = etree.parse(str(document_xml))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        body = tree.find("w:body", ns)
        if body is not None:
            for sect_pr in tree.xpath("//w:pPr/w:sectPr", namespaces=ns):
                sect_pr.getparent().remove(sect_pr)
            final_sect = body.find("w:sectPr", ns)
            if final_sect is not None:
                type_node = final_sect.find("w:type", ns)
                if type_node is not None:
                    final_sect.remove(type_node)
        tree.write(str(document_xml), xml_declaration=True, encoding="UTF-8", standalone="yes")

        temp_docx = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(temp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for path in tmp_path.rglob("*"):
                if path.is_file():
                    zout.write(path, path.relative_to(tmp_path).as_posix())
        temp_docx.replace(docx_path)


def convert_to_pdf(source: Path, out_dir: Path, final_name: str) -> Path:
    soffice = find_soffice()
    with tempfile.TemporaryDirectory(prefix="paper-polish-soffice-") as profile:
        run(
            [
                str(soffice),
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(source),
            ]
        )
    generated = out_dir / f"{source.stem}.pdf"
    final_pdf = out_dir / final_name
    if generated != final_pdf:
        generated.replace(final_pdf)
    return final_pdf


def convert_to_latex(source: Path, out_dir: Path, slug: str) -> Path:
    tex_path = out_dir / f"{slug}.tex"
    media_dir = out_dir / "media"
    media_dir.mkdir(parents=True, exist_ok=True)
    run(
        [
            "pandoc",
            str(source),
            "--from",
            "docx",
            "--to",
            "latex",
            "--standalone",
            "--extract-media",
            str(media_dir),
            "--output",
            str(tex_path),
        ]
    )
    text = tex_path.read_text(encoding="utf-8").replace(f"{media_dir.as_posix()}/", "media/")
    tex_path.write_text(text, encoding="utf-8")
    return tex_path


def write_review_notes() -> None:
    notes = """# Public Paper Edition Review Notes

These notes track defects and improvement targets identified while preparing public editions. The final report and applied focus editions preserve the original content and focus on formatting, spelling, and presentation cleanup.

## Final Report

- The title page is informative but sparse. The polished edition keeps the same metadata and improves spacing and hierarchy.
- Several section headings were stored as visually formatted normal paragraphs. The polished edition normalizes heading styles for easier navigation and cleaner PDF output.
- Large code-output screenshots are important evidence, but their surrounding captions need stronger spacing and hierarchy.
- The references spill onto a mostly empty final page in the original render. The polished edition tightens margins and paragraph rhythm to improve page economy.
- Technical claims around equilibrium and exploitability should eventually receive a deeper text review, but this formatting pass does not rewrite the research argument.

## Applied Focus Deliverable

- The title page is sparse and visually disconnected from the body. The polished edition improves type hierarchy while preserving the original metadata.
- Figures and diagrams carry much of the paper's argument, but spacing around them is inconsistent in the original render.
- Some all-caps headings were plain paragraphs rather than navigable heading styles. The polished edition normalizes those headings.
- The bibliography spills into a mostly empty final page in the original render. The polished edition improves page economy without removing sources.
- A later content pass should check source formatting consistency and tighten transitions, but this pass avoids substantive rewriting.

## Research Notebook

- The original notebook contains tab divider pages and section breaks that create mostly blank pages.
- Informal placeholders were present: "ASK ABOUT THIS" and "To fill." These were replaced with cleaner planning-note language.
- The abbreviation "envt" was expanded to "environment."
- The polished edition removes disruptive section breaks, applies consistent heading styles, and keeps the notebook as a research log rather than rewriting it as a formal paper.
"""
    (POLISHED_DIR / "review-notes.md").write_text(notes, encoding="utf-8")
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(POLISHED_DIR / "review-notes.md", PUBLIC_DIR / "review-notes.md")


def build() -> None:
    POLISHED_DIR.mkdir(parents=True, exist_ok=True)
    if PUBLIC_DIR.exists():
        shutil.rmtree(PUBLIC_DIR)
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)

    index = []
    for paper in PAPERS:
        if not paper.source.exists():
            raise FileNotFoundError(f"Missing source DOCX: {paper.source}")
        out_dir = POLISHED_DIR / paper.slug
        if out_dir.exists():
            shutil.rmtree(out_dir)
        out_dir.mkdir(parents=True)

        docx_path = out_dir / f"{paper.slug}.docx"
        shutil.copy2(paper.source, docx_path)
        document = Document(docx_path)
        apply_style_system(document, paper.mode)
        polish_paragraphs(document, paper.mode)
        document.save(docx_path)
        if paper.mode == "notebook":
            remove_notebook_section_breaks(docx_path)

        pdf_path = convert_to_pdf(docx_path, out_dir, f"{paper.slug}.pdf")
        tex_path = convert_to_latex(docx_path, out_dir, paper.slug)

        public_slug_dir = PUBLIC_DIR / paper.slug
        public_slug_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(docx_path, public_slug_dir / f"{paper.slug}.docx")
        shutil.copy2(tex_path, public_slug_dir / f"{paper.slug}.tex")
        media_dir = out_dir / "media"
        if media_dir.exists():
            shutil.copytree(media_dir, public_slug_dir / "media", dirs_exist_ok=True)
        shutil.copy2(pdf_path, PUBLIC_DIR / f"{paper.slug}.pdf")

        index.append(
            {
                "slug": paper.slug,
                "title": paper.title,
                "role": paper.role,
                "pdf": f"{paper.slug}.pdf",
                "docx": f"{paper.slug}/{paper.slug}.docx",
                "latex": f"{paper.slug}/{paper.slug}.tex",
            }
        )

    (POLISHED_DIR / "index.json").write_text(json.dumps(index, indent=2) + "\n", encoding="utf-8")
    (PUBLIC_DIR / "index.json").write_text(json.dumps(index, indent=2) + "\n", encoding="utf-8")
    write_review_notes()


if __name__ == "__main__":
    build()
